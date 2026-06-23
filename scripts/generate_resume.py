#!/usr/bin/env python3
"""
generate_resume.py — 把结构化简历内容渲染成统一专业模板的 docx。

复刻「实习.skill」案例使用的简历模板视觉规格：
- 字体 Aptos + 微软雅黑（现代无衬线），窄页边距（适合一页装下）
- 姓名大字号居中 + 联系方式行
- "求职意向"定位头段落
- section 标题带品牌色底部分隔线（如「教育背景」「核心经历」）
- 经历条目：标题加粗 + 时间右对齐，下挂 bullet

输入：一个 JSON 文件（结构见 resume_schema 注释 / scripts/sample_resume.json）
输出：docx 文件

用法：
    python3 generate_resume.py <input.json> [-o output.docx]
"""
import json
import argparse
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LATIN_FONT = "Aptos"
EAST_ASIA_FONT = "Microsoft YaHei"
NAME_SIZE = 18      # 姓名
SECTION_SIZE = 12   # section 标题
BODY_SIZE = 10      # 正文
RIGHT_TAB_CM = 17.0  # 时间右对齐 tab 位置（按窄边距 A4 正文宽度）
TEXT_COLOR = (34, 34, 34)
MUTED_COLOR = (90, 90, 90)
ACCENT_COLOR = (26, 83, 92)
ACCENT_HEX = "1A535C"
TRANSFER_LABEL_RE = re.compile(r"^\s*(?:[-•]\s*)?(?:迁移句|迁移说明|可迁移性)\s*[:：]\s*")


def warn(message):
    print(f"⚠️ {message}", file=sys.stderr)


def text_or_default(value, default):
    value = str(value).strip() if value is not None else ""
    return value or default


def as_list(value):
    if isinstance(value, list):
        return value
    if value in (None, ""):
        return []
    return [value]


def set_font(run, size=BODY_SIZE, bold=False, color=None, italic=False):
    run.font.name = LATIN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    else:
        run.font.color.rgb = RGBColor(*TEXT_COLOR)
    # 中文字体单独指定，避免中文被西文字体硬套。
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), EAST_ASIA_FONT)
    rfonts.set(qn('w:ascii'), LATIN_FONT)
    rfonts.set(qn('w:hAnsi'), LATIN_FONT)


def clean_bullet_text(text):
    """清理上游 prompt 泄漏出的标签词，避免简历里出现「迁移句:」。"""
    text = str(text or "").strip()
    return TRANSFER_LABEL_RE.sub("", text)


def normalize_resume_data(data):
    """把不完整 JSON 降级成可渲染结构，并把问题输出为 warning。"""
    if not isinstance(data, dict):
        raise ValueError("输入 JSON 顶层必须是 object。")

    normalized = {
        "name": text_or_default(data.get("name"), "姓名待补"),
        "contact": text_or_default(data.get("contact"), "联系方式待补"),
        "objective": text_or_default(data.get("objective"), ""),
        "sections": [],
    }
    if normalized["name"] == "姓名待补":
        warn("缺少 name，已使用占位文本。")
    if normalized["contact"] == "联系方式待补":
        warn("缺少 contact，已使用占位文本。")

    raw_sections = data.get("sections")
    if not isinstance(raw_sections, list):
        warn("缺少 sections 或 sections 不是数组，已生成“信息待补”占位区。")
        raw_sections = [
            {
                "title": "信息待补",
                "entries": [
                    {
                        "title": "请补充教育背景、核心经历、技能等内容",
                        "bullets": ["当前 JSON 缺少可渲染的 sections。"],
                    }
                ],
            }
        ]

    for section_index, section in enumerate(raw_sections, start=1):
        if not isinstance(section, dict):
            warn(f"第 {section_index} 个 section 不是 object，已跳过。")
            continue

        section_type = section.get("type")
        normalized_section = {
            "title": text_or_default(section.get("title"), f"未命名模块 {section_index}"),
        }

        if section_type == "skills":
            items = []
            for item_index, item in enumerate(as_list(section.get("items")), start=1):
                if isinstance(item, dict):
                    category = text_or_default(item.get("category"), f"技能类别 {item_index}")
                    content = text_or_default(item.get("content"), "")
                    if not content:
                        warn(f"技能区第 {item_index} 项缺少 content，已跳过。")
                        continue
                    items.append({"category": category, "content": content})
                else:
                    content = text_or_default(item, "")
                    if content:
                        items.append(content)
            if not items:
                warn(f"技能 section「{normalized_section['title']}」没有可渲染 items，已跳过。")
                continue
            normalized_section["type"] = "skills"
            normalized_section["items"] = items
        else:
            entries = []
            for entry_index, entry in enumerate(as_list(section.get("entries")), start=1):
                if not isinstance(entry, dict):
                    warn(f"section「{normalized_section['title']}」第 {entry_index} 条 entry 不是 object，已跳过。")
                    continue
                title = text_or_default(entry.get("title"), f"经历待补 {entry_index}")
                bullets = [clean_bullet_text(b) for b in as_list(entry.get("bullets"))]
                bullets = [b for b in bullets if b]
                entries.append(
                    {
                        "title": title,
                        "meta": text_or_default(entry.get("meta"), ""),
                        "subtitle": text_or_default(entry.get("subtitle"), ""),
                        "bullets": bullets,
                    }
                )
                if title.startswith("经历待补"):
                    warn(f"section「{normalized_section['title']}」第 {entry_index} 条缺少 title，已使用占位标题。")
            if not entries:
                warn(f"section「{normalized_section['title']}」没有可渲染 entries，已跳过。")
                continue
            normalized_section["entries"] = entries

        normalized["sections"].append(normalized_section)

    if not normalized["sections"]:
        warn("没有可渲染 section，已生成“信息待补”占位区。")
        normalized["sections"] = [
            {
                "title": "信息待补",
                "entries": [
                    {
                        "title": "请补充教育背景、核心经历、技能等内容",
                        "bullets": ["当前 JSON 没有可渲染内容。"],
                    }
                ],
            }
        ]

    return normalized


def add_bottom_border(paragraph):
    """给段落加底部分隔线（section 标题用）。"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), ACCENT_HEX)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def tight(paragraph, before=2, after=2):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.05


def add_name(doc, name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(p, 0, 2)
    set_font(p.add_run(name), NAME_SIZE, bold=True, color=ACCENT_COLOR)


def add_contact(doc, contact):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(p, 0, 6)
    set_font(p.add_run(contact), BODY_SIZE, color=MUTED_COLOR)


def add_objective(doc, objective):
    """求职意向定位头。"""
    p = doc.add_paragraph()
    tight(p, 2, 6)
    r = p.add_run("求职意向：")
    set_font(r, BODY_SIZE, bold=True, color=ACCENT_COLOR)
    set_font(p.add_run(objective), BODY_SIZE)


def add_section(doc, title):
    p = doc.add_paragraph()
    tight(p, 8, 3)
    set_font(p.add_run(title), SECTION_SIZE, bold=True, color=ACCENT_COLOR)
    add_bottom_border(p)


def add_entry(doc, title, meta=None, subtitle=None, bullets=None):
    """一条经历：标题(加粗) + 右对齐 meta(时间/地点)，可选副标题，下挂 bullets。"""
    p = doc.add_paragraph()
    tight(p, 4, 1)
    if meta:
        # 设置右对齐 tab
        pf = p.paragraph_format
        tab_stops = pf.tab_stops
        tab_stops.add_tab_stop(Cm(RIGHT_TAB_CM), WD_TAB_ALIGNMENT.RIGHT)
        set_font(p.add_run(title), BODY_SIZE, bold=True)
        set_font(p.add_run("\t" + meta), BODY_SIZE, color=MUTED_COLOR)
    else:
        set_font(p.add_run(title), BODY_SIZE, bold=True)

    if subtitle:
        sp = doc.add_paragraph()
        tight(sp, 0, 1)
        set_font(sp.add_run(subtitle), BODY_SIZE, color=MUTED_COLOR, italic=True)

    for b in (bullets or []):
        b = clean_bullet_text(b)
        if not b:
            continue
        bp = doc.add_paragraph(style=None)
        tight(bp, 0, 1)
        bp.paragraph_format.left_indent = Cm(0.5)
        set_font(bp.add_run("• " + b), BODY_SIZE)


def add_skills(doc, skills):
    """skills: list of {category, content} 或 list of str。"""
    for item in skills:
        p = doc.add_paragraph()
        tight(p, 1, 1)
        if isinstance(item, dict):
            set_font(p.add_run(item["category"] + "："), BODY_SIZE, bold=True)
            set_font(p.add_run(item["content"]), BODY_SIZE)
        else:
            set_font(p.add_run("• " + item), BODY_SIZE)


def build(data, out_path):
    data = normalize_resume_data(data)
    doc = Document()
    # 窄页边距
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.9)
        s.right_margin = Cm(1.9)
    # 默认样式字体
    normal = doc.styles['Normal']
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(BODY_SIZE)

    add_name(doc, data["name"])
    add_contact(doc, data["contact"])
    if data.get("objective"):
        add_objective(doc, data["objective"])

    for section in data["sections"]:
        add_section(doc, section["title"])
        if section.get("type") == "skills":
            add_skills(doc, section["items"])
        else:
            for e in section["entries"]:
                add_entry(
                    doc,
                    title=e["title"],
                    meta=e.get("meta"),
                    subtitle=e.get("subtitle"),
                    bullets=e.get("bullets"),
                )

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description="生成统一模板的 docx 简历")
    ap.add_argument("input", help="结构化简历 JSON 文件")
    ap.add_argument("-o", "--output", help="输出 docx 路径")
    args = ap.parse_args()

    with open(args.input, encoding="utf-8") as f:
        data = json.load(f)

    out = args.output or args.input.rsplit(".", 1)[0] + ".docx"
    build(data, out)
    print(f"✓ 已生成简历：{out}")


if __name__ == "__main__":
    main()
